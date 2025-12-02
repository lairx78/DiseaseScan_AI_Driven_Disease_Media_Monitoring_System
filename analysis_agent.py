result_dict = {}

import os
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "project78-456610-b3f1efc8ee09.json"

import getpass
import uuid
import json
import re
import requests
from typing import Annotated, Optional, Literal, List, Dict
from typing_extensions import TypedDict
from bs4 import BeautifulSoup
import traceback
from docx import Document

from langchain.chat_models import init_chat_model
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
import warnings
with warnings.catch_warnings():
    warnings.simplefilter("ignore")
    from langchain_core.pydantic_v1 import BaseModel, Field 
from langchain_core.tools import tool
from langchain_core.messages import AnyMessage, ToolMessage, HumanMessage, AIMessage
from langchain.globals import set_debug
from langgraph.graph import StateGraph, END, START
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode, tools_condition
from langchain_core.runnables import RunnableConfig, RunnableLambda
from langchain_core.chat_history import InMemoryChatMessageHistory

class AgentState(TypedDict):
    messages: Annotated[list[AnyMessage], add_messages]
    identified_input_type: Optional[Literal["url", "filename", "text", "unknown"]] = None
    identified_value: Optional[str] = None
    fetched_content_filename: Optional[str] = None
    error: Optional[str] = None

class InitialClassification(BaseModel):
    disease_related: Literal["yes", "no"] = Field(...)

class DiseaseNames(BaseModel):
    disease_names: Optional[List[str]] = Field(default=None)

class SummarizationResult(BaseModel):
    summary: Optional[str] = Field(default=None)

class ArticleTitle(BaseModel):
    title: Optional[str] = Field(default=None, description="The main title or headline of the article content.")

class ArticleAnalysis(BaseModel):
    disease_related: Literal["yes", "no"] = Field(...)
    article_title: Optional[str] = Field(default=None)
    disease_names: Optional[List[str]] = Field(default=None)
    summary: Optional[str] = Field(default=None)

class InputIdentificationResult(BaseModel):
    input_type: Literal["url", "filename", "text", "unknown"] = Field(...)
    value: Optional[str] = Field(default=None)

class FetchHtmlInput(BaseModel):
    url: str = Field(...)

class FetchHtmlOutput(BaseModel):
    success: bool = Field(...)
    message: str = Field(...)
    output_filename: Optional[str] = Field(default=None)


# --- Initialize Vertex AI LLMs ---
model_string_tool = "gemini-2.0-flash"
model_string = "gemini-2.5-flash-preview-04-17"
print(f"Using model: {model_string}")
print(f"Using model for tools: {model_string_tool}")
classification_llm_instance = init_chat_model(model=model_string_tool, model_provider="google_vertexai", temperature=0, convert_system_message_to_human=True)
extraction_llm_instance = init_chat_model(model=model_string_tool, model_provider="google_vertexai", temperature=0, convert_system_message_to_human=True)
summarization_llm_instance = init_chat_model(model=model_string_tool, model_provider="google_vertexai", temperature=0, convert_system_message_to_human=True)
agent_llm_instance = init_chat_model(model=model_string, model_provider="google_vertexai", temperature=0, convert_system_message_to_human=True)
print("Initialized LLMs.")

# --- Create the structured output runnables ---
structured_classifier = None
try:
    structured_classifier = classification_llm_instance.with_structured_output(InitialClassification)
    print("Using native .with_structured_output() with the classification LLM.")
except Exception as e:
     print(f"Critical error setting up structured classification LLM: {e}")
     structured_classifier = RunnableLambda(lambda _: InitialClassification(disease_related="no"))


# --- Prompts ---
classification_prompt_template = ChatPromptTemplate.from_template(
"""
Analyze the following text, which was extracted from an HTML page and may include navigation, headers, footers, and other non-article elements.
Your task is to first identify the **main article content** within the provided text.
Based **only** on the content of that main article, determine if it discusses communicable diseases. Communicable diseases are illnesses that can be spread from one person (or animal) to another (for example, influenza, tuberculosis, Ebola, etc.), whereas non-communicable diseases such as cancer, diabetes, and cardiovascular diseases do not spread between people.
Ignore any mentions of diseases found solely in navigation menus, sidebars, or footers.
Respond ONLY with a JSON object matching the 'InitialClassification' schema, with a key "disease_related" set to "yes" if the main article discusses communicable diseases or "no" otherwise.
Text:
{input}
"""
)
extraction_prompt_template = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            "You are an expert extraction algorithm. The provided text comes from a webpage and might include navigation, footers, etc.\n"
            "1. Identify the **main body/article content** within the text.\n"
            "2. Extract the names of any communicable diseases mentioned **exclusively within that main article content**.\n"
            "3. Ignore disease names found in surrounding website elements (menus, headers, footers, links).\n"
            "Return the names as a list of strings within the 'disease_names' field in the 'DiseaseNames' JSON schema.\n"
            "If no communicable diseases are mentioned *in the main article*, return null or an empty list for 'disease_names'."
        ),
        ("human", "Text:\n{text}"),
    ]
)
summarization_prompt_template = ChatPromptTemplate.from_messages(
    [
        (
            "system", 
            "You are an expert summarizer. You are given text extracted from an HTML page which may include extraneous website elements "
            "(such as navigation menus, sidebars, headers, footers, cookie banners, and legal text) as well as the main article content. \n\n"
            "Your task is to:\n"
            "1. Identify and extract the main article or body content from the provided text, ignoring non-essential website elements.\n"
            "2. Include key details such as the title, publication dates, main organizations, important statistics, and any recommendations or interventions described.\n"
            "3. Write a coherent, comprehensive yet concise summary in one or two paragraphs that covers all critical points of the main article content only.\n"
            "4. Do NOT include any information from non-article parts like navigation links, menus, sidebars, headers, footers, or cookie information.\n\n"
            "Respond ONLY with the JSON object matching the 'SummarizationResult' schema, containing the summary of the main article."
        ),
        ("human", "Article Text:\n{text}"),
    ]
)
title_extraction_prompt_template = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            "The following text was extracted from a webpage and may include navigation, headers, etc.\n"
            "1. Identify the **main article content** within the text.\n"
            "2. Determine the most prominent headline or title associated **specifically with that main article content**.\n"
            "3. Ignore titles from the overall webpage (<title> tag) if they don't match the main article's heading.\n"
            "Respond ONLY with the JSON object matching the 'ArticleTitle' schema, providing the title in the 'title' field.\n"
            "If no clear article title/headline can be discerned from the main content, return null for the 'title'."
        ),
        ("human", "Text:\n{text}"),
    ]
)


# --- Define the Tools ---
@tool
def identify_input_type(user_message: str) -> str:
    """
    Analyzes the user's message to determine if it is a URL, an HTML filename, or plain text.
    Use this tool FIRST on the user's latest message to decide the next step.
    Returns a JSON string ('InputIdentificationResult') indicating the type and value.
    Priority: url > filename > text.
    """
    print(f"\n--- Calling Identify Input Type Tool ---")
    print(f"Input to identifier: '{user_message}'")

    if not isinstance(user_message, str):
        result = InputIdentificationResult(input_type="unknown", value=None)
        return result.json()

    cleaned_message = user_message.strip()
    lower_message = cleaned_message.lower()

    if "http://" in lower_message or "https://" in lower_message:
        url_match = re.search(r'https?://\S+', cleaned_message, re.IGNORECASE)
        if url_match:
            result_type = "url"
            result_value = url_match.group()
            print(f"Identified as: url ({result_value})")
        else:
            result_type = "url"
            result_value = cleaned_message
            print(f"Identified as: url (default to full message)")   

    elif ".htm" in lower_message or ".html" in lower_message:
        match = re.search(r'[\w.-]+\.(?:htm|html)', cleaned_message, re.IGNORECASE)
        if match:
            result_type = "filename"
            result_value = match.group()
            print(f"Identified as: filename ({result_value})")
        else:
            result_type = "filename"
            result_value = cleaned_message
            print(f"Identified as: filename (default to full message)")

    elif cleaned_message:
        result_type = "text"
        result_value = cleaned_message
        print("Identified as: text")
    else:
        result_type = "unknown"
        result_value = None
        print("Identified as: unknown (empty input)")

    result = InputIdentificationResult(input_type=result_type, value=result_value)
    return result.json()

@tool(args_schema=FetchHtmlInput)
def fetch_and_save_html(url: str) -> str:
    """
    Fetches HTML from the provided URL and saves it locally to 'website_content.html'.
    Use this after 'identify_input_type' confirms the input is a URL.
    Returns a JSON string ('FetchHtmlOutput') indicating success/failure and the output filename if successful.
    """
    global result_dict

    print(f"\n--- Calling fetch_and_save_html Tool ---")
    print(f"Attempting to fetch URL: {url}")
    output_filename = "website_content.html"

    if not url or not isinstance(url, str) or not url.lower().startswith(('http://', 'https://')):
        error_msg = "Error: Invalid or missing URL. Please provide a full URL starting with 'http://' or 'https://'."
        print(error_msg)
        result = FetchHtmlOutput(success=False, message=error_msg, output_filename=None)
        result_dict = {}
        return result.json()

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }

    try:
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        html_content = response.text

        try:
            with open(output_filename, "w", encoding="utf-8") as f:
                f.write(html_content)
            success_msg = f"Successfully fetched HTML from {url} and saved it to '{output_filename}'."
            print(success_msg)
            result = FetchHtmlOutput(success=True, message=success_msg, output_filename=output_filename)
            return result.json()
        except OSError as e:
            error_msg = f"Error: Could not save file '{output_filename}'. Reason: {e}"
            print(error_msg)
            result = FetchHtmlOutput(success=False, message=error_msg, output_filename=None)
            result_dict = {}
            return result.json()

    except requests.exceptions.Timeout:
        error_msg = f"Error: Request timed out while trying to fetch {url}."
        print(error_msg)
        result = FetchHtmlOutput(success=False, message=error_msg, output_filename=None)
        result_dict = {}
        return result.json()
    except requests.exceptions.HTTPError as e:
        error_msg = f"Error: HTTP Error fetching {url}. Status: {e.response.status_code}. Reason: {e.response.reason}"
        print(error_msg)
        result = FetchHtmlOutput(success=False, message=error_msg, output_filename=None)
        result_dict = {}
        return result.json()
    except requests.exceptions.ConnectionError as e:
        error_msg = f"Error: Could not connect to {url}. Check URL/connection. Details: {e}"
        print(error_msg)
        result = FetchHtmlOutput(success=False, message=error_msg, output_filename=None)
        result_dict = {}
        return result.json()
    except requests.exceptions.RequestException as e:
        error_msg = f"Error: Request failed for {url}. Reason: {e}"
        print(error_msg)
        result = FetchHtmlOutput(success=False, message=error_msg, output_filename=None)
        result_dict = {}
        return result.json()
    except Exception as e:
        error_msg = f"Error: An unexpected error occurred during fetch/save for {url}. Reason: {e}"
        print(error_msg)
        result = FetchHtmlOutput(success=False, message=error_msg, output_filename=None)
        result_dict = {}
        return result.json()


@tool
def analyze_communicable_disease_article(input_content: str) -> str:
    """
    Analyzes input (text or local filename like 'file.htm', 'website_content.html').
    Reads file if filename provided. Determines if main content relates to communicable
    diseases, extracts article title, disease names, and summarizes if relevant.
    If the article is communicable disease related, the tool creates a Word report (.docx file)
    containing the details, then returns a message including the location of the file along with
    a plain text report.
    """

    global result_dict
    result_dict = {}

    print(f"\n--- Calling Analysis Tool ---")
    print(f"Input to analysis tool: '{input_content[:100]}...'")
    article_text = ""
    input_is_filename = False
    filename_provided = ""
    html_content_for_title = None

    if not isinstance(input_content, str) or not input_content.strip():
        return "Error: No valid text or filename provided to analysis tool."

    cleaned_input = input_content.strip()
    is_local_html = cleaned_input.lower().endswith((".htm", ".html"))
    is_temp_file = cleaned_input == "website_content.html"
    contains_path_sep = "/" in cleaned_input or "\\" in cleaned_input

    if is_temp_file or (is_local_html and not contains_path_sep):
        input_is_filename = True
        filename_provided = cleaned_input
        print(f"Analysis tool received filename: {filename_provided}. Attempting to read.")
        try:
            with open(filename_provided, 'r', encoding='utf-8') as file:
                html_content_for_title = file.read()
            print(f"Successfully read file: {filename_provided}")
            soup = BeautifulSoup(html_content_for_title, 'html.parser')
            article_text = soup.get_text(separator=' ', strip=True)
            print(f"Extracted text from HTML (first 200 chars): {article_text[:200]}...")
            if not article_text and not html_content_for_title.strip():
                return f"Error: File '{filename_provided}' was read but seems empty or contained no extractable text."
        except FileNotFoundError:
            print(f"Error: File not found by analysis tool: {filename_provided}")
            return f"Error: Analysis failed. The required file '{filename_provided}' was not found."
        except Exception as e:
            print(f"Error reading/parsing file {filename_provided} in analysis tool: {e}\n{traceback.format_exc()}")
            return f"Error: Could not read or parse the file '{filename_provided}'. Reason: {e}"
    else:
        print("Analysis tool received raw text.")
        article_text = input_content
        html_content_for_title = None

    if not article_text.strip() and not (input_is_filename and html_content_for_title):
        return "Error: Provided input resulted in empty content for analysis."
    if not structured_classifier or not extraction_llm_instance or not summarization_llm_instance:
        return "Error: One or more LLMs required for analysis were not initialized correctly."

    print(f"\n--- Calling Analysis Logic ---")
    final_analysis = ArticleAnalysis(disease_related="no")

    try:
        print("Step 1: Performing initial classification...")
        text_for_classification = article_text if article_text else ""
        classification_prompt = classification_prompt_template.invoke({"input": text_for_classification})
        classification_result: InitialClassification = structured_classifier.invoke(classification_prompt)
        final_analysis.disease_related = classification_result.disease_related
        print(f"Initial classification result: {final_analysis.disease_related}")

        result_dict["disease_related"] = final_analysis.disease_related

        if final_analysis.disease_related == "yes":
            text_for_details = article_text if article_text else ""
            content_for_title_extraction = html_content_for_title if html_content_for_title else text_for_details

            print("Step 2a: Extracting article title...")
            try:
                structured_title_extractor = extraction_llm_instance.with_structured_output(ArticleTitle)
                title_prompt = title_extraction_prompt_template.invoke({"text": content_for_title_extraction})
                title_result: ArticleTitle = structured_title_extractor.invoke(title_prompt)
                if title_result.title and title_result.title.strip():
                    final_analysis.article_title = title_result.title.strip()
                    print(f"Extracted article title: {final_analysis.article_title}")

                    result_dict["article_title"] = final_analysis.article_title
                else:
                    print("No specific article title extracted.")
            except Exception as e:
                print(f"Error during article title extraction: {e}\n{traceback.format_exc()}")

            print("Step 2b: Extracting disease names...")
            try:
                structured_extractor = extraction_llm_instance.with_structured_output(DiseaseNames)
                extraction_prompt = extraction_prompt_template.invoke({"text": text_for_details})
                extraction_result: DiseaseNames = structured_extractor.invoke(extraction_prompt)
                if extraction_result.disease_names:
                    final_analysis.disease_names = [str(name).strip() for name in extraction_result.disease_names if str(name).strip()]
                    print(f"Extracted disease names: {final_analysis.disease_names}")

                    result_dict["disease_names"] = final_analysis.disease_names
                else:
                    final_analysis.disease_names = []
                    print("No specific disease names extracted.")
            except Exception as e:
                print(f"Error during disease name extraction: {e}\n{traceback.format_exc()}")
                final_analysis.disease_names = ["Error during extraction"]

            print("Step 2c: Summarizing article...")
            try:
                structured_summarizer = summarization_llm_instance.with_structured_output(SummarizationResult)
                summarization_prompt = summarization_prompt_template.invoke({"text": text_for_details})
                summary_result: SummarizationResult = structured_summarizer.invoke(summarization_prompt)
                if summary_result.summary and summary_result.summary.strip():
                    final_analysis.summary = summary_result.summary.strip()
                    print(f"Generated summary: {final_analysis.summary[:100]}...")

                    result_dict["summary"] = final_analysis.summary
                else:
                    print("No summary generated.")
            except Exception as e:
                print(f"Error during summarization: {e}\n{traceback.format_exc()}")
                final_analysis.summary = "Error during summarization"

        source_info = f"File '{filename_provided}'" if input_is_filename else "Provided text"
        plain_report_lines = []
        plain_report_lines.append(f"Analysis of: {source_info}")
        if final_analysis.disease_related == "yes":
            plain_report_lines.append("The content IS related to communicable diseases.")
            if final_analysis.article_title:
                plain_report_lines.append(f"Article Title: {final_analysis.article_title}")
            else:
                plain_report_lines.append("Article Title: Not found.")
            if final_analysis.disease_names == ["Error during extraction"]:
                plain_report_lines.append("Disease Names: Could not reliably extract due to an error.")
            elif final_analysis.disease_names:
                plain_report_lines.append(f"Diseases Mentioned: {', '.join(final_analysis.disease_names)}.")
            else:
                plain_report_lines.append("Diseases Mentioned: None identified in the main article.")
            if final_analysis.summary == "Error during summarization":
                plain_report_lines.append("Summary: Could not be generated due to an error.")
            elif final_analysis.summary:
                plain_report_lines.append(f"Summary: {final_analysis.summary}")
            else:
                plain_report_lines.append("Summary: Not generated.")

            document = Document()
            document.add_heading("WORD REPORT", level=1)
            document.add_paragraph(f"Article Title: {final_analysis.article_title if final_analysis.article_title else 'Not found'}")
            if final_analysis.disease_names == ["Error during extraction"]:
                document.add_paragraph("Disease Names: Could not reliably extract due to an error.")
            elif final_analysis.disease_names:
                document.add_paragraph(f"Diseases Mentioned: {', '.join(final_analysis.disease_names)}")
            else:
                document.add_paragraph("Diseases Mentioned: None identified")
            if final_analysis.summary == "Error during summarization":
                document.add_paragraph("Summary: Could not be generated due to an error.")
            elif final_analysis.summary:
                document.add_heading("Summary:", level=2)
                document.add_paragraph(final_analysis.summary)
            else:
                document.add_paragraph("Summary: Not generated")

            report_filename = f"communicable_disease_report_{uuid.uuid4().hex}.docx"
            document.save(report_filename)

            combined_output = f"Word report generated: {report_filename}\n\n" + "\n".join(plain_report_lines)
            return combined_output
        else:
            plain_report_lines.append("The content is NOT related to communicable diseases.")
            return "\n".join(plain_report_lines)

    except Exception as e:
        print(f"Error during analysis tool execution (outer try block): {e}\n{traceback.format_exc()}")
        return f"An unexpected error occurred during the analysis process: {e}"


# --- Update Tool List ---
available_tools = [identify_input_type, fetch_and_save_html, analyze_communicable_disease_article]


# === Define the Agent ===
agent_prompt = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            "You are an article analysis assistant that specializes in analyzing content related to communicable diseases. You operate using three tools:\n"
            "1. `identify_input_type`: Determines if the user's input is a URL, a filename, or text. ALWAYS use this tool first when new input is provided.\n"
            "2. `fetch_and_save_html`: Fetches content from a URL and saves it as 'website_content.html'. Use ONLY if `identify_input_type` returns 'url'.\n"
            "3. `analyze_communicable_disease_article`: Analyzes text or a local file (including 'website_content.html') for communicable disease content. Use AFTER step 1 (or step 2 if URL) returns text or a filename. This tool returns a full analysis report.\n\n"
            "Your process is as follows:\n\n"
            "A. New Input:\n"
            "   - When a new input (text, URL, or filename) is provided, your ONLY first action is to call `identify_input_type` using the user's exact message.\n\n"
            "B. Tool Processing:\n"
            "   - Based on the result from `identify_input_type`:\n"
            "       • If it returns 'url' with a URL, then call `fetch_and_save_html` with that URL.\n"
            "       • If it returns 'filename' with a filename, then call `analyze_communicable_disease_article` using that filename.\n"
            "       • If it returns 'text' with text, then call `analyze_communicable_disease_article` using that text.\n"
            "       • If it returns 'unknown', prompt the user for valid input (text, URL, or a .htm/.html filename).\n"
            "   - If `fetch_and_save_html` returns success with an 'output_filename', then call `analyze_communicable_disease_article` with that filename (i.e. 'website_content.html').\n"
            "   - If `fetch_and_save_html` fails, report the error message from that tool.\n"
            "   - When `analyze_communicable_disease_article` completes, output its result verbatim.\n\n"
            "C. Follow-Up Questions (Without New Input):\n"
            "   - For any follow-up question, first determine if it is related to the previously analyzed article. A question is considered related if it explicitly or implicitly refers to the article (for example, it uses phrases such as 'the article', 'previous analysis', 'as mentioned in the report', or it refers directly to details extracted from the article such as the mentioned disease).\n"
            "   - Next, determine if the question is health-related. A question is considered health-related if it asks for an explanation, further details, or elaboration on a disease or any other health aspect extracted from the article.\n"
            "   - If and only if BOTH of these conditions are met (i.e. the follow-up is related to the analyzed article and is health-related), then provide an answer using only the extracted information. If the article does not include additional details on the subject, state that it does not supply further information.\n"
            "   - In all other cases — either because the question is not related to the article or because it is related but not health-related — respond with:\n"
            "     \"My analysis is strictly limited to the content provided in the article. Please ask questions related to that content or provide new input for analysis.\"\n\n"
            "Remember: Do not provide external or general information beyond what is contained in the analyzed article."
            "--- EXAMPLES OF CORRECT BEHAVIOR ---\n\n"
            "**Example 1: URL Input (Communicable Disease Found)**\n\n"
            "Human: Please analyze this page: https://www.who.int/emergencies/disease-outbreak-news/item/1996_01_22c-en\n"
            "AI: (Calls `identify_input_type` with 'https://www.who.int/emergencies/disease-outbreak-news/item/1996_01_22c-en')\n"
            "Tool (`identify_input_type` result): {'input_type': 'url', 'value': 'https://www.who.int/emergencies/disease-outbreak-news/item/1996_01_22c-en'}\n"
            "AI: (Calls `fetch_and_save_html` with 'https://www.who.int/emergencies/disease-outbreak-news/item/1996_01_22c-en')\n"
            "Tool (`fetch_and_save_html` result): {'success': true, 'message': 'Successfully fetched...', 'output_filename': 'website_content.html'}\n"
            "AI: (Calls `analyze_communicable_disease_article` with 'website_content.html')\n"
            "Tool (`analyze_communicable_disease_article` result): Analysis of: File 'website_content.html'\n"
            "The content IS related to communicable diseases.\n"
            "Article Title: 1996 - Liberia\n"
            "Diseases Mentioned: Ebola virus.\n"
            "Summary: In January 1996, a suspected case of Ebola virus was reported to WHO. A 25-year-old male from Liberia was admitted... WHO does not recommend travel restrictions.\n"
            "AI: Analysis of: File 'website_content.html'\n"
            "The content IS related to communicable diseases.\n"
            "Article Title: 1996 - Liberia\n"
            "Diseases Mentioned: Ebola virus.\n"
            "Summary: In January 1996, a suspected case of Ebola virus was reported to WHO. A 25-year-old male from Liberia was admitted... WHO does not recommend travel restrictions.\n\n"
            "**Example 2: Text Input (Not Communicable Disease)**\n\n"
            "Human: Check this text: Local elections concluded yesterday with high turnout.\n"
            "AI: (Calls `identify_input_type` with 'Local elections concluded yesterday with high turnout.')\n"
            "Tool (`identify_input_type` result): {'input_type': 'text', 'value': 'Local elections concluded yesterday with high turnout.'}\n"
            "AI: (Calls `analyze_communicable_disease_article` with 'Local elections concluded yesterday with high turnout.')\n"
            "Tool (`analyze_communicable_disease_article` result): Analysis of: Text input\n"
            "The article is NOT related to communicable diseases.\n"
            "AI: Analysis of: Text input\n"
            "The article is NOT related to communicable diseases.\n\n"
            "**Example 3: Fetch Error**\n\n"
            "Human: Analyze https://thissitedoesnotexistreally123.com\n"
            "AI: (Calls `identify_input_type` with 'https://thissitedoesnotexistreally123.com')\n"
            "Tool (`identify_input_type` result): {'input_type': 'url', 'value': 'https://thissitedoesnotexistreally123.com'}\n"
            "AI: (Calls `fetch_and_save_html` with 'https://thissitedoesnotexistreally123.com')\n"
            "Tool (`fetch_and_save_html` result): {'success': false, 'message': 'Error: Could not connect to https://thissitedoesnotexistreally123.com...'}\n"
            "AI: Error: Could not connect to https://thissitedoesnotexistreally123.com...\n"
            "--- END EXAMPLES ---"
        ),
        MessagesPlaceholder(variable_name="messages", optional=True),
    ]
)


agent_runnable = agent_llm_instance.bind_tools(available_tools)

def agent_node(state: AgentState, config: RunnableConfig):
    print("\n--- Calling Agent Node ---")
    last_message = state["messages"][-1] if state["messages"] else None
    if isinstance(last_message, ToolMessage):
        print(f"Agent sees ToolMessage from tool call ID related to: {last_message.tool_call_id}")
    result = agent_runnable.invoke(state['messages'], config) # Pass the whole message list
    return {"messages": [result]}


# === Define the Graph ===
def handle_tool_error(state: AgentState) -> dict:
    error_msg = state.get("error", "Unknown error")
    print(f"--- Handling Tool Error: {error_msg} ---")
    last_message = state["messages"][-1] if state["messages"] else None
    tool_calls = getattr(last_message, 'tool_calls', []) if isinstance(last_message, AIMessage) else []
    if not isinstance(tool_calls, list): tool_calls = []

    messages = []
    failed_tool_call_id = None
    if tool_calls:
        failed_tool_call_id = tool_calls[0].get("id") if isinstance(tool_calls[0], dict) else None

    if failed_tool_call_id:
         messages = [
            ToolMessage(
                content=f"Error executing tool: {error_msg}. Please check input or configuration.",
                tool_call_id=failed_tool_call_id,
            )
        ]
    else:
        messages.append(AIMessage(content=f"An error occurred processing the request: {error_msg}"))

    return {"messages": messages, "identified_input_type": None, "identified_value": None, "fetched_content_filename": None, "error": None}


tool_node = ToolNode(available_tools).with_fallbacks(
    [RunnableLambda(handle_tool_error)], exception_key="error"
)

builder = StateGraph(AgentState)
builder.add_node("agent", agent_node)
builder.add_node("tools", tool_node)
builder.add_edge(START, "agent")
builder.add_conditional_edges(
    "agent",
    tools_condition,
    {"tools": "tools", END: END},
)
builder.add_edge("tools", "agent") 
graph = builder.compile()

# === Use the Graph ===
_printed_ids = set() 
conversation_histories = {}  

def run_conversation(input_message: str, thread_id: str):
    config = {"configurable": {"thread_id": thread_id}}
    print(f"\nUser: {input_message}")
    global _printed_ids, conversation_histories
    history_messages = conversation_histories.get(thread_id, [])
    history_messages.append(HumanMessage(content=input_message))
    state = {"messages": history_messages}
    try:
        events = graph.stream(state, config, stream_mode="values")
        last_event = None
        print("\n--- Graph Execution Start ---")
        for i, event in enumerate(events):
            print(f"\nStream Event {i}: {list(event.keys())}")
            if "messages" in event:
                current_message = event["messages"][-1]
                if current_message.id not in _printed_ids:
                    print(f"\n{type(current_message).__name__} (ID: {current_message.id})")
                    print(f"CONTENT: {getattr(current_message, 'content', 'N/A')}")
                    if isinstance(current_message, AIMessage) and hasattr(current_message, 'tool_calls') and current_message.tool_calls:
                        print(f"TOOL CALLS: {current_message.tool_calls}")
                    if isinstance(current_message, ToolMessage):
                        print(f"TOOL CALL ID: {getattr(current_message, 'tool_call_id', 'N/A')}")
                    _printed_ids.add(current_message.id)
            last_event = event
        last_message = last_event["messages"][-1] if last_event and "messages" in last_event and last_event["messages"] else None
        conversation_histories[thread_id] = last_event["messages"] if last_event and "messages" in last_event else history_messages
    except Exception as e:
        print(f"\n--- Error during graph execution: {e} ---")
        print(traceback.format_exc())
        last_message = AIMessage(content=f"An unexpected error occurred: {e}")
    finally:
        print("\n--- Conversation Turn Finished ---")
    return last_message